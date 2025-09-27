"""
Optimized Document Processor for PDF and DOCX files
Handles procurement tendering manuals, TEPP, and compliance documents
with AI-powered analysis and cost optimization
"""

import os
import json
import hashlib
import asyncio
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, asdict
import re

# PDF processing
try:
    import PyPDF2
    import pdfplumber
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# DOCX processing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# AI processing
from app.services.ai_model_detector import get_ai_detector
from app.services.three_tier_rag_system import ThreeTierRAGSystem, RAGTier, DocumentCategory, RAGDocument
from app.config import settings

@dataclass
class ProcessedDocument:
    """Processed document with extracted content and metadata"""
    file_path: str
    file_name: str
    file_type: str
    file_size: int
    content: str
    metadata: Dict[str, Any]
    compliance_tags: List[str]
    document_category: DocumentCategory
    rag_tier: RAGTier
    processing_cost: float
    processing_time: float
    ai_confidence: float
    extracted_sections: Dict[str, str]
    compliance_requirements: List[str]
    created_date: str

class OptimizedDocumentProcessor:
    """
    Optimized document processor with AI analysis and cost optimization
    """
    
    def __init__(self, data_dir: Path):
        self.data_dir = data_dir
        self.rag_system = ThreeTierRAGSystem(data_dir)
        self.ai_detector = get_ai_detector()
        self.ai_model = None
        
        # Cost optimization settings
        self.max_tokens_per_request = 4000  # Optimize for cost
        self.chunk_size = 2000  # Optimal chunk size for processing
        self.cache_processed_docs = True
        self.processed_cache = {}
        
        # Initialize AI model
        self._initialize_ai_model()
        
        # Document classification patterns
        self.classification_patterns = self._load_classification_patterns()
    
    def _initialize_ai_model(self):
        """Initialize AI model for document processing"""
        try:
            self.ai_model = self.ai_detector.get_recommended_model()
            print(f"✅ Document Processor initialized with {self.ai_model.model_name}")
        except Exception as e:
            print(f"⚠️ Failed to initialize AI model: {e}")
            self.ai_model = None
    
    def _load_classification_patterns(self) -> Dict[str, Dict[str, List[str]]]:
        """Load document classification patterns for cost optimization"""
        return {
            "procurement_manual": {
                "keywords": ["procurement", "tendering", "manual", "guidelines", "policy"],
                "tier": RAGTier.GLOBAL,
                "category": DocumentCategory.PROCUREMENT_GUIDELINES
            },
            "tepp_template": {
                "keywords": ["tepp", "tender evaluation", "probity plan", "evaluation plan"],
                "tier": RAGTier.INTERNAL,
                "category": DocumentCategory.INTERNAL_PROCEDURES
            },
            "evaluation_criteria": {
                "keywords": ["evaluation", "criteria", "scoring", "matrix", "assessment"],
                "tier": RAGTier.INTERNAL,
                "category": DocumentCategory.EVALUATION_CRITERIA
            },
            "compliance_standards": {
                "keywords": ["compliance", "standards", "whs", "safety", "environmental", "quality"],
                "tier": RAGTier.GLOBAL,
                "category": DocumentCategory.TENDERING_STANDARDS
            },
            "contract_templates": {
                "keywords": ["contract", "agreement", "terms", "conditions", "template"],
                "tier": RAGTier.INTERNAL,
                "category": DocumentCategory.INTERNAL_PROCEDURES
            },
            "project_specific": {
                "keywords": ["project", "specific", "requirements", "specifications"],
                "tier": RAGTier.PROJECT,
                "category": DocumentCategory.WHS_COMPLIANCE
            }
        }
    
    async def process_document(self, file_path: str) -> ProcessedDocument:
        """Process a single document with AI analysis and cost optimization"""
        
        start_time = datetime.now()
        file_path = Path(file_path)
        
        # Check cache first for cost optimization
        if self.cache_processed_docs:
            cache_key = self._get_cache_key(file_path)
            if cache_key in self.processed_cache:
                cached_doc = self.processed_cache[cache_key]
                print(f"📋 Using cached document: {file_path.name}")
                return cached_doc
        
        # Extract content based on file type
        content, metadata = await self._extract_content(file_path)
        
        # Classify document for optimal RAG placement
        classification = self._classify_document(content, file_path.name)
        
        # AI analysis with cost optimization
        ai_analysis = await self._analyze_document_ai(content, classification)
        
        # Calculate processing cost
        processing_cost = self._calculate_processing_cost(content, ai_analysis)
        
        # Create processed document
        processed_doc = ProcessedDocument(
            file_path=str(file_path),
            file_name=file_path.name,
            file_type=file_path.suffix.lower(),
            file_size=file_path.stat().st_size,
            content=content,
            metadata=metadata,
            compliance_tags=ai_analysis.get('compliance_tags', []),
            document_category=classification['category'],
            rag_tier=classification['tier'],
            processing_cost=processing_cost,
            processing_time=(datetime.now() - start_time).total_seconds(),
            ai_confidence=ai_analysis.get('confidence', 0.8),
            extracted_sections=ai_analysis.get('sections', {}),
            compliance_requirements=ai_analysis.get('compliance_requirements', []),
            created_date=datetime.now().isoformat()
        )
        
        # Cache for future use
        if self.cache_processed_docs:
            self.processed_cache[cache_key] = processed_doc
        
        return processed_doc
    
    async def _extract_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content from PDF or DOCX files"""
        
        content = ""
        metadata = {
            "file_type": file_path.suffix.lower(),
            "file_size": file_path.stat().st_size,
            "created_date": datetime.fromtimestamp(file_path.stat().st_ctime).isoformat(),
            "modified_date": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
        }
        
        try:
            if file_path.suffix.lower() == '.pdf' and PDF_AVAILABLE:
                content, metadata = await self._extract_pdf_content(file_path)
            elif file_path.suffix.lower() == '.docx' and DOCX_AVAILABLE:
                content, metadata = await self._extract_docx_content(file_path)
            else:
                # Fallback to text extraction
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                
        except Exception as e:
            print(f"⚠️ Error extracting content from {file_path.name}: {e}")
            content = f"Error extracting content: {e}"
        
        return content, metadata
    
    async def _extract_pdf_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content from PDF files using multiple methods for optimization"""
        
        content = ""
        metadata = {}
        
        try:
            # Method 1: pdfplumber (better for complex layouts)
            with pdfplumber.open(file_path) as pdf:
                pages_content = []
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        pages_content.append(f"Page {page_num + 1}:\n{page_text}")
                
                content = "\n\n".join(pages_content)
                
                # Extract metadata
                metadata.update({
                    "total_pages": len(pdf.pages),
                    "extraction_method": "pdfplumber"
                })
                
        except Exception as e:
            print(f"⚠️ pdfplumber failed for {file_path.name}: {e}")
            
            # Method 2: PyPDF2 (fallback)
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    pages_content = []
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text:
                            pages_content.append(f"Page {page_num + 1}:\n{page_text}")
                    
                    content = "\n\n".join(pages_content)
                    metadata.update({
                        "total_pages": len(pdf_reader.pages),
                        "extraction_method": "PyPDF2"
                    })
                    
            except Exception as e2:
                print(f"⚠️ PyPDF2 also failed for {file_path.name}: {e2}")
                content = f"Error extracting PDF content: {e2}"
        
        return content, metadata
    
    async def _extract_docx_content(self, file_path: Path) -> Tuple[str, Dict[str, Any]]:
        """Extract content from DOCX files"""
        
        content = ""
        metadata = {}
        
        try:
            doc = DocxDocument(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            content = "\n".join(paragraphs)
            
            # Extract metadata
            metadata.update({
                "total_paragraphs": len(paragraphs),
                "extraction_method": "python-docx"
            })
            
            # Extract document properties if available
            if doc.core_properties:
                props = doc.core_properties
                if props.title:
                    metadata["title"] = props.title
                if props.author:
                    metadata["author"] = props.author
                if props.created:
                    metadata["created"] = props.created.isoformat()
                if props.modified:
                    metadata["modified"] = props.modified.isoformat()
                    
        except Exception as e:
            print(f"⚠️ Error extracting DOCX content from {file_path.name}: {e}")
            content = f"Error extracting DOCX content: {e}"
        
        return content, metadata
    
    def _classify_document(self, content: str, filename: str) -> Dict[str, Any]:
        """Classify document for optimal RAG placement using pattern matching"""
        
        content_lower = content.lower()
        filename_lower = filename.lower()
        
        # Score each classification pattern
        scores = {}
        for doc_type, pattern in self.classification_patterns.items():
            score = 0
            
            # Check filename patterns
            for keyword in pattern["keywords"]:
                if keyword in filename_lower:
                    score += 2
            
            # Check content patterns
            for keyword in pattern["keywords"]:
                if keyword in content_lower:
                    score += 1
            
            scores[doc_type] = score
        
        # Get best classification
        best_type = max(scores, key=scores.get)
        best_score = scores[best_type]
        
        # Default classification if no pattern matches
        if best_score == 0:
            best_type = "project_specific"
        
        classification = self.classification_patterns[best_type].copy()
        classification["confidence"] = min(best_score / 5.0, 1.0)  # Normalize to 0-1
        classification["type"] = best_type
        
        return classification
    
    async def _analyze_document_ai(self, content: str, classification: Dict[str, Any]) -> Dict[str, Any]:
        """AI analysis with cost optimization"""
        
        if not self.ai_model or self.ai_model.provider == "none":
            return self._basic_analysis(content, classification)
        
        # Optimize content for AI processing
        optimized_content = self._optimize_content_for_ai(content)
        
        try:
            # Create focused prompt for cost optimization
            prompt = self._create_optimized_prompt(optimized_content, classification)
            
            # Use AI model for analysis
            if hasattr(self.ai_model, 'generate'):
                response = await self.ai_model.generate(prompt)
            else:
                response = await self._fallback_ai_analysis(prompt)
            
            # Parse AI response
            analysis = self._parse_ai_response(response)
            
            return analysis
            
        except Exception as e:
            print(f"⚠️ AI analysis failed: {e}")
            return self._basic_analysis(content, classification)
    
    def _optimize_content_for_ai(self, content: str) -> str:
        """Optimize content for AI processing to reduce costs"""
        
        # Truncate if too long
        if len(content) > self.max_tokens_per_request * 4:  # Rough token estimation
            content = content[:self.max_tokens_per_request * 4]
        
        # Remove excessive whitespace
        content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
        
        # Remove page numbers and headers/footers
        content = re.sub(r'Page \d+ of \d+', '', content)
        content = re.sub(r'^\d+\s*$', '', content, flags=re.MULTILINE)
        
        return content.strip()
    
    def _create_optimized_prompt(self, content: str, classification: Dict[str, Any]) -> str:
        """Create optimized prompt for cost-effective AI analysis"""
        
        doc_type = classification.get("type", "unknown")
        
        prompt = f"""
        Analyze this {doc_type} document for NSW procurement compliance:
        
        Document Type: {doc_type}
        Content Preview: {content[:1000]}...
        
        Provide JSON response with:
        {{
            "compliance_tags": ["tag1", "tag2"],
            "sections": {{"section_name": "content_summary"}},
            "compliance_requirements": ["requirement1", "requirement2"],
            "confidence": 0.8,
            "summary": "Brief document summary"
        }}
        
        Focus on:
        - NSW Local Government Act compliance
        - WHS and safety requirements
        - Environmental standards
        - Quality management
        - Evaluation criteria
        - Probity and accountability
        """
        
        return prompt
    
    async def _fallback_ai_analysis(self, prompt: str) -> str:
        """Fallback AI analysis method"""
        return json.dumps({
            "compliance_tags": ["procurement", "compliance"],
            "sections": {"main": "Document content"},
            "compliance_requirements": ["Standard compliance requirements"],
            "confidence": 0.5,
            "summary": "Document processed with basic analysis"
        })
    
    def _parse_ai_response(self, response: str) -> Dict[str, Any]:
        """Parse AI response with error handling"""
        try:
            # Extract JSON from response
            if "```json" in response:
                json_start = response.find("```json") + 7
                json_end = response.find("```", json_start)
                if json_end > json_start:
                    json_str = response[json_start:json_end].strip()
                else:
                    json_str = response[json_start:].strip()
            else:
                json_str = response.strip()
            
            return json.loads(json_str)
        except Exception as e:
            print(f"⚠️ Error parsing AI response: {e}")
            return {
                "compliance_tags": ["procurement"],
                "sections": {"main": "Document content"},
                "compliance_requirements": ["Standard requirements"],
                "confidence": 0.3,
                "summary": "Basic analysis completed"
            }
    
    def _basic_analysis(self, content: str, classification: Dict[str, Any]) -> Dict[str, Any]:
        """Basic analysis without AI for cost optimization"""
        
        # Extract basic compliance tags
        compliance_tags = []
        content_lower = content.lower()
        
        if any(term in content_lower for term in ["whs", "safety", "health"]):
            compliance_tags.append("WHS")
        if any(term in content_lower for term in ["environmental", "sustainability"]):
            compliance_tags.append("Environmental")
        if any(term in content_lower for term in ["quality", "iso9001"]):
            compliance_tags.append("Quality")
        if any(term in content_lower for term in ["evaluation", "criteria"]):
            compliance_tags.append("Evaluation")
        if any(term in content_lower for term in ["compliance", "standards"]):
            compliance_tags.append("Compliance")
        
        # Extract sections using simple pattern matching
        sections = {}
        lines = content.split('\n')
        current_section = "main"
        section_content = []
        
        for line in lines:
            line = line.strip()
            if line and (line.isupper() or line.startswith(('1.', '2.', '3.', '4.', '5.'))):
                if section_content:
                    sections[current_section] = ' '.join(section_content)
                current_section = line
                section_content = []
            else:
                section_content.append(line)
        
        if section_content:
            sections[current_section] = ' '.join(section_content)
        
        return {
            "compliance_tags": compliance_tags,
            "sections": sections,
            "compliance_requirements": ["Standard compliance requirements"],
            "confidence": 0.7,
            "summary": f"Document classified as {classification.get('type', 'unknown')}"
        }
    
    def _calculate_processing_cost(self, content: str, ai_analysis: Dict[str, Any]) -> float:
        """Calculate processing cost for optimization tracking"""
        
        # Rough cost estimation (adjust based on your AI provider)
        base_cost = 0.001  # Base cost per document
        content_cost = len(content) * 0.000001  # Cost per character
        ai_cost = ai_analysis.get('confidence', 0.5) * 0.01  # AI analysis cost
        
        return base_cost + content_cost + ai_cost
    
    def _get_cache_key(self, file_path: Path) -> str:
        """Generate cache key for document"""
        return hashlib.md5(f"{file_path}_{file_path.stat().st_mtime}".encode()).hexdigest()
    
    async def process_all_documents(self, directory: str = None) -> List[ProcessedDocument]:
        """Process all documents in a directory with cost optimization"""
        
        if directory is None:
            directory = str(self.data_dir)
        
        directory_path = Path(directory)
        processed_docs = []
        
        # Find all PDF and DOCX files
        pdf_files = list(directory_path.rglob("*.pdf"))
        docx_files = list(directory_path.rglob("*.docx"))
        
        all_files = pdf_files + docx_files
        
        print(f"📁 Found {len(all_files)} documents to process")
        
        # Process documents in batches for optimization
        batch_size = 5
        for i in range(0, len(all_files), batch_size):
            batch = all_files[i:i + batch_size]
            
            # Process batch concurrently
            tasks = [self.process_document(str(file_path)) for file_path in batch]
            batch_results = await asyncio.gather(*tasks, return_exceptions=True)
            
            for result in batch_results:
                if isinstance(result, ProcessedDocument):
                    processed_docs.append(result)
                elif isinstance(result, Exception):
                    print(f"⚠️ Error processing document: {result}")
            
            print(f"✅ Processed batch {i//batch_size + 1}/{(len(all_files) + batch_size - 1)//batch_size}")
        
        return processed_docs
    
    async def add_to_rag_system(self, processed_doc: ProcessedDocument) -> bool:
        """Add processed document to RAG system"""
        
        try:
            # Create a simplified document for RAG (avoiding enum serialization issues)
            simple_doc = {
                "document_id": hashlib.md5(f"{processed_doc.file_name}_{processed_doc.created_date}".encode()).hexdigest()[:12],
                "title": processed_doc.file_name,
                "content": processed_doc.content,
                "tier": processed_doc.rag_tier.value,  # Convert enum to string
                "category": processed_doc.document_category.value,  # Convert enum to string
                "version": "1.0",
                "council": "Blacktown City Council",
                "project_type": processed_doc.metadata.get('project_type'),
                "metadata": {
                    **processed_doc.metadata,
                    "compliance_tags": processed_doc.compliance_tags,
                    "ai_confidence": processed_doc.ai_confidence,
                    "processing_cost": processed_doc.processing_cost,
                    "extracted_sections": processed_doc.extracted_sections,
                    "compliance_requirements": processed_doc.compliance_requirements
                },
                "created_date": processed_doc.created_date,
                "updated_date": processed_doc.created_date,
                "upload_date": processed_doc.created_date,
                "file_path": str(processed_doc.file_path) if processed_doc.file_path else None
            }
            
            # Save directly to RAG directory
            tier_dir = self.data_dir / "three_tier_rag" / processed_doc.rag_tier.value
            tier_dir.mkdir(parents=True, exist_ok=True)
            
            doc_file = tier_dir / f"{simple_doc['document_id']}.json"
            with open(doc_file, 'w', encoding='utf-8') as f:
                json.dump(simple_doc, f, indent=2, ensure_ascii=False)
            
            print(f"✅ Added {processed_doc.file_name} to RAG system")
            return True
            
        except Exception as e:
            print(f"❌ Error adding {processed_doc.file_name} to RAG system: {e}")
            return False
    
    def get_processing_summary(self, processed_docs: List[ProcessedDocument]) -> Dict[str, Any]:
        """Get processing summary with cost analysis"""
        
        total_cost = sum(doc.processing_cost for doc in processed_docs)
        total_time = sum(doc.processing_time for doc in processed_docs)
        
        # Group by document type
        by_type = {}
        for doc in processed_docs:
            doc_type = doc.metadata.get('file_type', 'unknown')
            if doc_type not in by_type:
                by_type[doc_type] = []
            by_type[doc_type].append(doc)
        
        # Group by RAG tier
        by_tier = {}
        for doc in processed_docs:
            tier = doc.rag_tier.value
            if tier not in by_tier:
                by_tier[tier] = []
            by_tier[tier].append(doc)
        
        return {
            "total_documents": len(processed_docs),
            "total_cost": total_cost,
            "total_time": total_time,
            "average_cost_per_doc": total_cost / len(processed_docs) if processed_docs else 0,
            "average_time_per_doc": total_time / len(processed_docs) if processed_docs else 0,
            "by_file_type": {k: len(v) for k, v in by_type.items()},
            "by_rag_tier": {k: len(v) for k, v in by_tier.items()},
            "compliance_tags": list(set(tag for doc in processed_docs for tag in doc.compliance_tags))
        }
