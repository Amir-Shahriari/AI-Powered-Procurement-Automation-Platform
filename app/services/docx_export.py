"""
Utility for exporting structured JSON documents to Word (DOCX).

This module writes DOCX directly with python-docx (no pandoc). It:
- Normalises escaped newlines (\r\n, \n, \\n, /n) into real line breaks
- Writes one paragraph per line
- Detects bullet-like prefixes (•, -, *, '1.', 'a.') and applies 'List Bullet'
- Renders lists of primitives as bullets
- Renders lists of dicts as tables
"""
from __future__ import annotations
from docx.shared import Inches

import json
import re
import tempfile
from pathlib import Path
from typing import Any, Dict, Iterable, List

from docx import Document


# -----------------------------
# Newline / bullet normalisation
# -----------------------------
_NEWLINE_RE = re.compile(r'(?:\r\n|\r|\n|\\n|/n)+')

_BULLET_PREFIX_RE = re.compile(
    r'^\s*(?:[\u2022•\-\*]|\d+\.|[a-zA-Z]\.)[\s\t]+'
)


def _normalize_to_lines(value: str | Iterable[str]) -> list[str]:
    """Return a clean list of non-empty lines from a string or list of strings."""
    lines: list[str] = []
    if isinstance(value, (list, tuple)):
        for item in value:
            s = str(item or "")
            s = _NEWLINE_RE.sub('\n', s)
            for ln in s.split('\n'):
                ln = ln.strip()
                if ln:
                    lines.append(ln)
        return lines

    s = str(value or "")
    s = _NEWLINE_RE.sub('\n', s)
    for ln in s.split('\n'):
        ln = ln.strip()
        if ln:
            lines.append(ln)
    return lines


def _strip_bullet_prefix(text: str) -> str:
    return _BULLET_PREFIX_RE.sub('', text, count=1).strip()


def write_block(doc: Document, value: Any, *, force_bullets: bool = False) -> None:
    """
    Write a text block to the doc:
      - Splits on newlines, handles escaped newlines
      - If lines look like bullets (or force_bullets=True), write each line as 'List Bullet'
      - Else, write each line as its own normal paragraph
    """
    # For non-strings/lists, just string-ify
    if not isinstance(value, (str, list, tuple)):
        value = "" if value is None else str(value)

    lines = _normalize_to_lines(value)
    if not lines:
        return

    # Decide bullet mode: either forced or any line appears bulleted
    bullet_mode = force_bullets or any(_BULLET_PREFIX_RE.match(ln) for ln in lines)

    for ln in lines:
        text = _strip_bullet_prefix(ln) if bullet_mode else ln
        p = doc.add_paragraph(text)
        if bullet_mode:
            try:
                p.style = 'List Bullet'  # Built-in Word style
            except Exception:
                # If style missing, keep as normal paragraph
                pass


# -----------------------------
# Table rendering helpers
# -----------------------------
def _is_list_of_dicts(items: list[Any]) -> bool:
    return bool(items) and all(isinstance(x, dict) for x in items if x is not None)


def _union_keys(items: list[dict]) -> list[str]:
    keys: list[str] = []
    for item in items:
        for k in item.keys():
            if k not in keys:
                keys.append(k)
    return keys


def _add_table(doc: Document, rows: list[dict], keys: list[str]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(keys))
    table.style = 'Light List Accent 1' if 'Light List Accent 1' in [s.name for s in doc.styles] else table.style

    # Header
    hdr = table.rows[0].cells
    for j, k in enumerate(keys):
        hdr[j].text = str(k)

    # Body
    for i, row in enumerate(rows, start=1):
        cells = table.rows[i].cells
        for j, k in enumerate(keys):
            v = row.get(k, "")
            if isinstance(v, (dict, list, tuple)):
                # Pretty-print nested structures
                cells[j].text = json.dumps(v, ensure_ascii=False, indent=2)
            else:
                cells[j].text = "" if v is None else str(v)


# -----------------------------
# Recursive writer
# -----------------------------
def _write_obj(doc: Document, obj: Any, *, level: int = 1, heading: str | None = None) -> None:
    """
    Recursively write JSON-like objects into the DOCX document.
    - Dicts → heading + recurse
    - Lists of dicts → table
    - Lists of primitives → bullet list
    - Primitives → paragraphs (with newline/bullet handling)
    """
    if heading:
        # python-docx headings: level 0..9 (0 is Title)
        doc.add_heading(heading, level=min(max(level, 1), 9))

    if isinstance(obj, dict):
        for key, value in obj.items():
            _write_obj(doc, value, level=level + 1, heading=str(key))
        return

    if isinstance(obj, list):
        if _is_list_of_dicts(obj):
            keys = _union_keys([x for x in obj if isinstance(x, dict)])
            _add_table(doc, [x for x in obj if isinstance(x, dict)], keys)
        else:
            # List of primitives → bullets
            write_block(doc, obj, force_bullets=True)
        return

    # Primitive value
    write_block(doc, obj, force_bullets=False)


# -----------------------------
# Public API
# -----------------------------
def export_json_to_docx(data: Dict[str, Any], title: str) -> str:
    """
    Generate a DOCX file from the given JSON data using python-docx.

    Args:
        data: Structured JSON-like object.
        title: Title for the document (written as Heading level 0).

    Returns:
        Path to the generated DOCX file.
    """
    doc = Document()

    # ---- Cover logos (first page) ----
    # Resolve logo paths (same convention as the UI)
    from pathlib import Path as _Path
    # app/app/services -> app/app
    app_root = _Path(__file__).resolve().parents[1]
    default_dir = app_root / "ui" / "assets"

    logo1 = _Path(os.getenv("APP_LOGO1")) if os.getenv("APP_LOGO1") else (default_dir / "lgp-logo-retina.png")
    logo2 = _Path(os.getenv("APP_LOGO2")) if os.getenv("APP_LOGO2") else (default_dir / "BlacktownCityCouncil_2019.png")

    # Place them in a 1x2 table to keep them on the same row
    try:
        tbl = doc.add_table(rows=1, cols=2)
        row = tbl.rows[0]
        if logo1.exists():
            row.cells[0].paragraphs[0].add_run().add_picture(str(logo1), width=Inches(1.2))
        if logo2.exists():
            row.cells[1].paragraphs[0].add_run().add_picture(str(logo2), width=Inches(1.2))
    except Exception:
        # If images are missing or python-docx can't load them, just skip logos gracefully
        pass

    # Title follows after the logos
    doc.add_heading(title, level=0)

    # Title
    doc.add_heading(title, level=0)

    # Body
    if isinstance(data, dict):
        for key, value in data.items():
            _write_obj(doc, value, level=1, heading=str(key))
    else:
        _write_obj(doc, data, level=1, heading=None)

    # Persist to a temp .docx
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        out_path = tmp.name
    doc.save(out_path)
    return out_path
